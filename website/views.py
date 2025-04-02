from collections import defaultdict
import uuid
from docx import Document
from flask import Blueprint, app, render_template, request, flash, jsonify, redirect, send_file, url_for
from flask_login import login_required, current_user
from .models import QuestionImage, Subject, Question, User, CourseOutcome
from . import db
from werkzeug.security import generate_password_hash
from io import BytesIO
from datetime import datetime
from itertools import combinations
from random import shuffle
from weasyprint import HTML
from docx import Document 
import os
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


views = Blueprint('views', __name__)

@views.route('/admin_dashboard', methods=['GET', 'POST'])
@login_required
def admin_dashboard():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('auth.login'))

    if request.method == 'POST':
        action = request.form.get('action') 

        if action == 'add-teacher':
            return redirect(url_for('views.add_teacher'))

        elif action == 'add-subject':
            return redirect(url_for('views.add_subject'))
        
        elif action == 'allocate-subject':
            return redirect(url_for('views.allocate_subject'))
        
        elif action == 'view-teachers':
            return redirect(url_for('views.view_teachers'))
        
        elif action == 'view-subjects':
            return redirect(url_for('views.view_subjects'))
                
        print(f"Action received: {action}")

    department_subjects = []
    if current_user.is_mod:
        # Query subjects based on the moderator's department
        department_subjects = Subject.query.filter_by(department=current_user.department).all()
   
    subjects = Subject.query.all()  # Fetch all subjects for admin to view
    return render_template("admin/admin_dashboard.html", user=current_user, subjects=subjects,department_subjects=department_subjects)


@views.route('/view_teachers', methods=['GET', 'POST'])
@login_required
def view_teachers():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('views.admin_dashboard'))
        
    # Query all non-admin users (teachers)
    teachers = User.query.filter(User.id != current_user.id).order_by(User.is_admin.desc()).all()

    return render_template('admin/view_teachers.html', users=teachers, user=current_user)


@views.route('/view_subjects', methods=['GET', 'POST'])
@login_required
def view_subjects():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('views.admin_dashboard'))
    
    subjects = Subject.query.all()
    return render_template('admin/view_subjects.html', subs=subjects, user=current_user)


@views.route('/add_teacher', methods=['GET', 'POST'])
@login_required
def add_teacher():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('views.admin_dashboard'))

    if request.method == 'POST':
        # Debugging: Print the form data to check what is being submitted
        print(request.form)

        teacher_name = request.form.get('name')
        email = request.form.get('email')
        password = request.form.get('password1')
        confirm_password = request.form.get('password2')
        designation = request.form.get('designation')
        department = request.form.get('dept')
        experience = request.form.get('exp')
        is_admin = request.form.get('is_admin')
        is_mod = request.form.get('is_mod') 

        if not teacher_name or not email or not password or not confirm_password or not designation or not department or not experience:
            flash('All fields are required!', category='error')
        elif password != confirm_password:
            flash('Passwords do not match!', category='error')
        else:
            try:
                # Convert `is_admin`and 'is_mod to a Boolean
                is_admin = True if is_admin == 'True' else False
                is_mod = True if is_mod == 'True' else False

                # Hash the password before storing it
                hashed_password = generate_password_hash(password, method='pbkdf2:sha256')

                # Create a new user (teacher)
                new_teacher = User(
                    email=email,
                    password=hashed_password,
                    firstname=teacher_name,
                    designation=designation,
                    department=department,
                    experience=int(experience),
                    is_admin=is_admin,
                    is_mod=is_mod
                )

                db.session.add(new_teacher)
                db.session.commit()

                flash('Teacher added successfully!', category='success')
            except Exception as e:
                db.session.rollback()
                flash(f'Error: {str(e)}', category='error')

    return render_template("admin/add_teacher.html", user=current_user)


@views.route('/add_subject', methods=['GET', 'POST'])
@login_required
def add_subject():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('auth.login'))

    if request.method == 'POST':
        # Retrieve CO data from the form
        co_numbers = request.form.getlist('co_number[]')
        co_descriptions = request.form.getlist('co_description[]')

        # Validate form inputs
        subject_code = request.form.get('code')
        subject_name = request.form.get('name')
        semester = request.form.get('sem')
        department = request.form.get('dept')
        academic_year = request.form.get('year')
        regulations = request.form.get('reg')

        if not subject_code or not subject_name or not semester or not department or not academic_year or not regulations or not co_numbers or not co_descriptions:
            flash('All fields are required, including COs!', category='error')
        else:
            try:
                # Create a new subject
                new_subject = Subject(
                    code=subject_code,
                    name=subject_name,
                    semester=int(semester),
                    department=department,
                    academic_year=int(academic_year),
                    regulations=regulations,
                    created_by=current_user.id
                )

                db.session.add(new_subject)
                db.session.commit()

                # Add Course Outcomes to the newly created subject
                for i in range(len(co_numbers)):
                    new_co = CourseOutcome(
                        subject_id=new_subject.id,
                        co_number=co_numbers[i],
                        description=co_descriptions[i]
                    )
                    db.session.add(new_co)

                db.session.commit()

                flash('Subject and Course Outcomes added successfully!', category='success')

            except Exception as e:
                db.session.rollback()
                flash(f'Error: {str(e)}', category='error')

    return render_template("admin/add_subject.html", user=current_user)



@views.route('/remove-subject/<int:subject_id>/', methods=['GET', 'POST'])
@login_required
def remove_subject(subject_id):
    # Fetch the subject to delete
    subject = Subject.query.get(subject_id)
    if subject:
        db.session.delete(subject)
        db.session.commit()
        # Optionally flash a success message
        flash("Subject removed successfully.", "success")
    else:
        # Flash an error message if the subject doesn't exist
        flash("Subject not found.", "error")
    return redirect(url_for('views.view_subjects'))


@views.route('/remove-user/<int:user_id>/', methods=['GET', 'POST'])
@login_required
def remove_user(user_id):
    # Fetch the subject to delete
    user = User.query.get(user_id)
    if user:
        db.session.delete(user)
        db.session.commit()
        # Optionally flash a success message
        flash("User removed successfully.", "success")
    else:
        # Flash an error message if the subject doesn't exist
        flash("User not found.", "error")
    return redirect(url_for('views.view_teachers'))



@views.route('/allocate_subject', methods=['GET', 'POST'])
@login_required
def allocate_subject():
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('auth.login'))

    teachers = User.query.all()  # Fetch all teachers
    subjects = Subject.query.all()  # Fetch all subjects

    print("Teachers:", teachers)
    print("Subjects:", subjects)

    return render_template(
        "admin/allocate_subject.html",
        user=current_user,
        teachers=teachers,
        subjects=subjects
    )


@views.route('/edit_subject/<int:subject_id>', methods=['GET', 'POST'])
@login_required
def edit_subject(subject_id):
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('auth.login'))

    # Fetch the subject by ID
    sub = Subject.query.get_or_404(subject_id)
    course_outcomes = CourseOutcome.query.filter_by(subject_id=sub.id).all()

    if request.method == 'POST':
        print(request.form)  # Debugging
        action = request.form.get('action')  # Get the value of the action button

        if action == 'edit-subject':
            # Check if all keys are present
            missing_keys = [key for key in ['code', 'name', 'semester', 'department', 'academic_year', 'regulations'] if key not in request.form]
            if missing_keys:
                flash(f'Missing fields: {", ".join(missing_keys)}', category='error')
                return redirect(url_for('views.edit_subject', subject_id=subject_id))

            # Update subject details
            sub.code = request.form.get('code')
            sub.name = request.form.get('name')
            sub.semester = request.form.get('semester')
            sub.department = request.form.get('department')
            sub.academic_year = request.form.get('academic_year')
            sub.regulations = request.form.get('regulations')

            for co in course_outcomes:
                co_number = request.form.getlist('co_number[]')
                co_description = request.form.getlist('co_description[]')
                remove_co = request.form.getlist('remove_co[]')

                # Update description if not marked for removal
                if str(co.id) not in remove_co:
                    co.description = co_description[co_number.index(co.co_number)]
                else:
                    db.session.delete(co)

            # Handle new COs
            new_co_numbers = request.form.getlist('new_co_number[]')
            new_co_descriptions = request.form.getlist('new_co_description[]')
            for number, description in zip(new_co_numbers, new_co_descriptions):
                if description.strip():
                    new_co = CourseOutcome(
                        subject_id=subject_id,
                        co_number=number,
                        description=description.strip()
                    )
                    db.session.add(new_co)

            db.session.commit()
            flash("Subject and COs updated successfully!", "success")
            return redirect(url_for('views.view_subjects'))

        return render_template('admin/edit_subject.html', subject=sub, course_outcomes=course_outcomes, user=current_user)





@views.route('/edit_user/<int:user1_id>', methods=['GET', 'POST'])
@login_required
def edit_user(user1_id):
    if not current_user.is_admin:
        flash('Access denied. Admins only.', category='error')
        return redirect(url_for('auth.login'))

    # Fetch the user by ID
    user1 = User.query.get_or_404(user1_id)
    if request.method == 'POST':
        print(request.form)  # Debugging
        action = request.form.get('action')  # Get the value of the action button

        if action == 'edit-user':
            # Check if all keys are present
            missing_keys = [key for key in ['name', 'email', 'designation', 'department', 'experience'] if key not in request.form]
            if missing_keys:
                flash(f'Missing fields: {", ".join(missing_keys)}', category='error')
                return redirect(url_for('views.edit_user', user1_id=user1_id))

            # Update subject details
            user1.name = request.form.get('name')
            user1.email = request.form.get('email')
            user1.designation = request.form.get('designation')
            user1.department = request.form.get('department')
            user1.experience = request.form.get('experience')
            is_admin = request.form.get('is_admin')
            is_admin = True if is_admin == 'True' else False
            user1.is_admin = is_admin
            is_mod = request.form.get('is_mod')
            is_mod = True if is_mod == 'True' else False
            user1.is_mod = is_mod

            db.session.commit()
            flash('Teacher details updated successfully!', category='success')
            return redirect(url_for('views.view_teachers'))

    return render_template('admin/edit_teacher.html', user=current_user, user1=user1)




@views.route('/assign_subject', methods=['POST'])
@login_required
def assign_subject():
    if not current_user.is_admin:
        return jsonify({"error": "Admins only."}), 403

    data = request.get_json()
    teacher_id = data.get('teacher_id')
    subject_id = data.get('subject_id')

    teacher = User.query.get(teacher_id)
    subject = Subject.query.get(subject_id)

    if teacher and subject:
        if subject not in teacher.subjects:
            teacher.subjects.append(subject)
            db.session.commit()
            return jsonify({"message": "Subject assigned successfully!"}), 200
        return jsonify({"error": "Subject already assigned."}), 400
    return jsonify({"error": "Invalid teacher or subject."}), 404


@views.route('/unassign_subject', methods=['POST'])
@login_required
def unassign_subject():
    if not current_user.is_admin:
        return jsonify({"error": "Admins only."}), 403

    data = request.get_json()
    teacher_id = data.get('teacher_id')
    subject_id = data.get('subject_id')

    teacher = User.query.get(teacher_id)
    subject = Subject.query.get(subject_id)

    if teacher and subject:
        if subject in teacher.subjects:
            teacher.subjects.remove(subject)
            db.session.commit()
            return jsonify({"message": "Subject unassigned successfully!"}), 200
        return jsonify({"error": "Subject not assigned to teacher."}), 400
    return jsonify({"error": "Invalid teacher or subject."}), 404




@views.route('/teacher_dashboard', methods=['GET', 'POST'])
@login_required
def teacher_dashboard():
    department_subjects = []
    if current_user.is_mod:
        # Query subjects based on the moderator's department
        department_subjects = Subject.query.filter_by(department=current_user.department).all()
    return render_template('staff/teacher_dashboard.html', user=current_user,department_subjects=department_subjects)


@views.route('/mod_questions/<int:subject_id>', methods=['GET', 'POST'])
@login_required
def mod_questions(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    subject_id = subject.id
    course_outcomes = CourseOutcome.query.filter_by(subject_id=subject_id).all()
    
    # Check if the user is assigned to the subject
    if subject.code not in [s.code for s in current_user.subjects] and not current_user.is_admin:
        flash("You are not authorized to view this subject's questions.", "danger")
        if current_user.is_admin:
            return redirect(url_for('views.admin_dashboard'))
        else:
            return redirect(url_for('views.teacher_dashboard'))

    # Handle the form submission for verify/unverify
    if request.method == 'POST':
        question_id = request.form.get('question_id')
        action = request.form.get('action')
        
        # Fetch the question from the database
        question = Question.query.get(question_id)
        if not question:
            flash("Question not found.", "error")
        else:
            # Update the verification status based on the action
            if action == 'verify':
                question.is_verified = True
                flash("Question successfully verified.", "success")
            elif action == 'unverify':
                question.is_verified = False
                flash("Question successfully unverified.", "success")
            else:
                flash("Invalid action.", "error")

            # Commit the changes to the database
            try:
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash("An error occurred while updating the question.", "error")

        # Redirect to the same page after processing
        return redirect(request.referrer or url_for('views.mod_questions', subject_id=subject.id))

    # Retrieve all questions for the subject, ordered by module_no
    questions = Question.query.filter_by(subject_code=subject.code).order_by(Question.module_no).all()
    
    # Group questions by module_no manually
    grouped_questions = defaultdict(list)
    for question in questions:
        grouped_questions[question.module_no].append(question)

    return render_template('mod/mod_questions.html', subject=subject, questions=grouped_questions, course_outcomes=course_outcomes, user=current_user)



@views.route('/add_questions/<int:subject_id>', methods=['GET', 'POST'])
@login_required
def add_questions(subject_id):
    # Fetch the subject details based on the subject_id
    subject = Subject.query.get_or_404(subject_id)
    course_outcomes = CourseOutcome.query.filter_by(subject_id=subject_id).all()

    # Check if the user is assigned to the subject
    if subject.code not in [s.code for s in current_user.subjects] and not current_user.is_admin:
        flash("You are not authorized to add this subject's questions.", "danger")
        if current_user.is_admin:
            return redirect(url_for('views.admin_dashboard'))
        else:
            return redirect(url_for('views.teacher_dashboard'))

    # Get the upload folder path from config
    upload_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'uploads'))
    print(upload_folder)  # Print the upload folder to check the path

    # If the user is submitting the form
    if request.method == 'POST':
        module_no = request.form.get('module_no')
        content = request.form.get('content')
        marks = request.form.get('marks')
        rbt_level = request.form.get('rbt_level')
        co = request.form.get('co')
        image_filenames = []

        # Step 1: Create the Question first to get its ID
        new_question = Question(
            subject_code=subject.code,
            module_no=module_no,
            content=content,
            marks=marks,
            rbt_level=rbt_level,
            co_id=co,
            created_by=current_user.id
        )
        db.session.add(new_question)
        db.session.commit()  # Commit to generate the Question ID (new_question.id)

        # Step 2: Process images and save them with the required filename format
        if 'images' in request.files:
            print("Got image/s")
            images = request.files.getlist('images')
            for image in images:
                if image and allowed_file(image.filename):
                    try:
                        # Generate the filename using subject name and question ID
                        extension = os.path.splitext(image.filename)[1]  # Get the file extension (e.g., .png, .jpg)
                        filename = f"{subject.name.replace(' ', '_')}-{new_question.id}{extension}"
                        filename = secure_filename(filename)

                        # Save the image
                        image_path = os.path.join(upload_folder, filename)
                        image.save(image_path)
                        image_filenames.append(filename)
                    except Exception as e:
                        flash(f"Error saving image: {e}", "danger")
                        return redirect(url_for('views.add_questions', subject_id=subject_id))

        # Step 3: Add image file paths to the QuestionImage model
        for filename in image_filenames:
            new_image = QuestionImage(
                question_id=new_question.id,
                filename=filename
            )
            db.session.add(new_image)

        db.session.commit()

        # Flash a success message
        flash(f"Question added successfully for {subject.name} (Code: {subject.code})!", category='success')

        # Stay on the same page to add another question
        return redirect(url_for('views.add_questions', subject_id=subject_id))


    # Fetch all questions already added for this subject (optional)
    questions = Question.query.filter_by(subject_code=subject.code).all()

    # Render the Add Question page
    return render_template('add_questions.html', user=current_user, subject=subject, course_outcomes=course_outcomes)



from flask import send_from_directory
@login_required
@views.route('/uploads/<path:filename>')
def uploaded_file(filename):
    # Serve files from the 'uploads' directory in your project root
    uploads_dir = os.path.join(os.path.abspath(os.path.dirname(__file__)), '..', 'uploads')
    return send_from_directory(uploads_dir, filename)




@views.route('/view_questions/<int:subject_id>', methods=['GET'])
@login_required
def view_questions(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    
    # Check if the user is assigned to the subject
    if subject.code not in [s.code for s in current_user.subjects] and not current_user.is_admin:
        flash("You are not authorized to view this subject's questions.", "danger")
        if current_user.is_admin:
            return redirect(url_for('views.admin_dashboard'))
        else:
            return redirect(url_for('views.teacher_dashboard'))

    # Retrieve all questions for the subject, ordered by module_no
    questions = Question.query.filter_by(subject_code=subject.code).order_by(Question.module_no).all()
    # Fetch images associated with each question
    images = QuestionImage.query.join(Question).filter(Question.subject_code == subject.code).all()
    
    # Group questions by module_no manually
    grouped_questions = defaultdict(list)
    for question in questions:
        grouped_questions[question.module_no].append(question)
        
    return render_template('view_questions.html', subject=subject, questions=grouped_questions, user=current_user, images=images)


@views.route('/edit_question/<int:question_id>', methods=['GET', 'POST'])
@login_required
def edit_question(question_id):
    question = Question.query.get_or_404(question_id)
    subject = Subject.query.filter_by(code=question.subject_code).first_or_404()
    subject_id = subject.id
    course_outcomes = CourseOutcome.query.filter_by(subject_id=subject_id).all()
    images = QuestionImage.query.filter_by(question_id=question_id).all()

    # Get the upload folder path from config
    upload_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'uploads'))
    print(upload_folder)  # Print the upload folder to check the path
    
    # Check if the user is authorised
    if subject.code not in [s.code for s in current_user.subjects] and not current_user.is_admin:
        flash("You are not authorized to edit this subject's questions.", "danger")
        if current_user.is_admin:
            return redirect(url_for('views.admin_dashboard'))
        else:
            return redirect(url_for('views.teacher_dashboard'))
    
    if request.method == 'POST':
        # Update the question fields with the form data
        question.module_no = request.form['module_no']
        question.content = request.form['content']
        question.marks = request.form['marks']
        question.rbt_level = request.form['rbt_level']
        question.co_id = request.form['co']
        new_image_filenames = []

        # Handle image removal
        remove_img_ids = request.form.getlist('remove_img[]')
        if remove_img_ids:
            for img_id in remove_img_ids:
                image = QuestionImage.query.get(img_id)
                if image:
                    db.session.delete(image)

        # Handle image addition
        if 'new_images' in request.files:
            new_images = request.files.getlist('new_images')
            for new_image in new_images:
                if new_image and allowed_file(new_image.filename):
                    try:
                        # Generate a unique filename
                        extension = os.path.splitext(new_image.filename)[1]
                        unique_id = str(uuid.uuid4())  # Generate a unique identifier
                        filename = f"{subject.name.replace(' ', '_')}-{question_id}-{unique_id}{extension}"
                        filename = secure_filename(filename)

                        # Save the image
                        new_image_path = os.path.join(upload_folder, filename)
                        new_image.save(new_image_path)
                        new_image_filenames.append(filename)
                    except Exception as e:
                        flash(f"Error saving image: {e}", "danger")
                        return redirect(url_for('views.edit_question', question_id=question_id))

        # Add image file paths to the QuestionImage model
        for filename in new_image_filenames:
            new_image = QuestionImage(
                question_id=question_id,
                filename=filename
            )
            db.session.add(new_image)

        db.session.commit()
        flash("Question updated successfully!", "success")
        return redirect(url_for('views.view_questions', subject_id=subject.id))
    
    return render_template('edit_question.html', question=question, user=current_user, subject=subject, course_outcomes=course_outcomes, images=images)



@views.route('/remove_question/<int:question_id>', methods=['POST'])
@login_required
def remove_question(question_id):
    question = Question.query.get_or_404(question_id)
    subject = Subject.query.filter_by(code=question.subject_code).first_or_404()
    
    # Check if the current user is authorized to remove this question
    if question.subject_code not in [subject.code for subject in current_user.subjects] and not current_user.is_admin:
        flash("You are not authorized to remove this question.", "danger")
        return redirect(url_for('views.view_questions', subject_id=subject.id))
    
    # Delete associated images from the database
    images = QuestionImage.query.filter_by(question_id=question.id).all()
    for image in images:
        # Construct the path to the image file
        image_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), '..', 'uploads', image.filename)
        
        # Delete the image file from the file system
        try:
            if os.path.exists(image_path):
                os.remove(image_path)  # Remove the image file
        except Exception as e:
            flash(f"Error deleting image file: {e}", "danger")
    
        # Delete the image record from the database
        db.session.delete(image)
    
    # Commit the image deletions
    db.session.commit()
    
    # Delete the question
    db.session.delete(question)
    db.session.commit()
    
    flash("Question removed successfully!", "success")
    return redirect(url_for('views.view_questions', subject_id=subject.id))


@views.route('/profile', methods=['GET'])
@login_required
def profile():
    return render_template('profile.html', user=current_user)



def find_valid_combinations(questions, target_marks, count):
    # Filter questions with marks <= target_marks
    filtered_questions = [q for q in questions if q.marks <= target_marks]

    # Shuffle the questions to introduce randomness
    shuffle(filtered_questions)

    # Generate combinations of questions with size `count`
    for combination in combinations(filtered_questions, count):
        if sum(q.marks for q in combination) == target_marks:
            return list(combination)  # Return the first valid combination found

    return None  # No valid combination found




@views.route('/auto_generate_qp', methods=['GET', 'POST'])
@login_required
def auto_generate_qp():
    if request.method == 'POST':
        # Get the output format from the submitted button
        output_format = request.form.get('output_format', 'pdf')  # Default to 'pdf'

        # Retrieve other form data
        subject_code = request.form.get('subject')
        exam_type = request.form.get('exam_type')
        academic_year = request.form.get('academic_year')
        exam_date = request.form.get('exam_date')
        max_marks = request.form.get('max_marks')
        start_time = request.form.get('start_time')
        duration = request.form.get('duration')  # Duration in hours
        modules = request.form.getlist('modules[]')
        num_questions = request.form.getlist('num_questions[]')
        module_max_marks = request.form.getlist('module_max_marks[]')
        additional_details = request.form.get('additional_details')

        # Validate and process the modules/questions
        total_module_marks = sum(int(marks) for marks in module_max_marks)
        max_marks = int(max_marks)  # Convert max_marks to integer

        if total_module_marks != max_marks:
            flash(f"Total marks for all modules ({total_module_marks}) do not match the maximum marks ({max_marks}).", category='error')
            return redirect(url_for('views.auto_generate_qp'))

        # Fetch subject details
        subject = Subject.query.filter_by(code=subject_code).first()
        if not subject:
            flash('Invalid subject selected.', category='error')
            return redirect(url_for('views.auto_generate_qp'))

        department = subject.department  # Get department of the subject

        exam_date = datetime.strptime(exam_date, "%Y-%m-%d").strftime("%d-%m-%Y")

        # Fetch questions for each module
        questions_by_module = {}

        for module, count, module_marks in zip(modules, num_questions, module_max_marks):
            count = int(count)
            module_marks = int(module_marks)

            module_questions = Question.query.filter_by(subject_code=subject_code, module_no=module).all()

            if len(module_questions) < count:
                flash(f"Not enough questions available for Module {module}.", category='error')
                return redirect(url_for('views.auto_generate_qp'))

            valid_combination = find_valid_combinations(module_questions, module_marks, count)
            if not valid_combination:
                flash(f"Could not find {count} questions for Module {module} that sum to {module_marks} marks.", category='error')
                return redirect(url_for('views.auto_generate_qp'))

            questions_by_module[module] = valid_combination

        if not any(questions_by_module.values()):
            flash("No questions available for the selected configuration.", category='error')
            return redirect(url_for('views.auto_generate_qp'))

        html = render_template(
            'question_paper.html',
            exam_type=exam_type,
            academic_year=academic_year,
            exam_date=exam_date,
            max_marks=max_marks,
            start_time=start_time,
            duration=duration,
            subject=subject,
            department=department,
            modules=modules,
            questions_by_module=questions_by_module,
            additional_details=additional_details
        )

        # Generate output based on the selected format
        if output_format == 'pdf':
            # Convert the HTML to a PDF
            pdf = HTML(string=html, base_url=request.url_root).write_pdf()

            # Create a BytesIO object for the binary content
            response = BytesIO()
            response.write(pdf)
            response.seek(0)

            return send_file(
                response,
                download_name="QuestionPaper.pdf",
                as_attachment=True
            )

        elif output_format == 'word':
            # Create Document
            doc = Document()

            image_path = os.path.join(os.getcwd(), 'website', 'static', 'images', 'uni.png')
            doc.add_picture(image_path, width=Inches(6))
            if doc.paragraphs:
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            heading = doc.add_paragraph('Department of Computer Science and Engineering')
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(14)

            # Add Exam Info
            doc.add_paragraph(f"{exam_type} – {academic_year}", style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph(f"Course Code: {subject.code} \t\t Maximum Marks: {'max_marks'}")
            doc.add_paragraph(f"Date: {'exam_date'} \t\t Time: {'start_time'} - Duration: {'duration'} hours")

            # Add Note
            note = doc.add_paragraph('NOTE: ')
            note.add_run('Answer All Questions').bold = True

            # Create Table for Questions
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            table.columns[0].width = Cm(0.001)  
            table.columns[1].width = Cm(20) 
            table.columns[2].width = Cm(0.001)  
            table.columns[3].width = Cm(0.001)  
            table.columns[4].width = Cm(0.001) 

            headers = ['Q#', 'Question', 'Marks', 'BL', 'CO']
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].runs[0].bold = True

            # Add Questions to Table
            question_number = 1
            for module, questions in questions_by_module.items():
                for question in questions:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(question_number)  # Q#
                    row_cells[1].text = question.content  # Question
                    row_cells[2].text = str(question.marks)  # Marks
                    row_cells[3].text = question.rbt_level  # BL (Bloom's Level)
                    row_cells[4].text = str(question.co_id)  # CO (Course Outcome)
                    question_number += 1
                               
            # Prepared and Approved Section
            doc.add_paragraph("\nPrepared By: Dr Loganathan R\t\tApproved By: The Principal")

            # Save to Memory Buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Send as Response
            return send_file(
                buffer,
                download_name=f"{subject.name}_QuestionPaper.docx",
                as_attachment=True,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    # For GET request, render the form
    current_year = datetime.now().year
    subjects = Subject.query.all()

    return render_template('admin/auto_generate_qp.html',
                           subjects=subjects,
                           current_year=current_year,
                           user=current_user)


@views.route('/manual_generate_qp', methods=['GET', 'POST'])
@login_required
def manual_generate_qp():
    # Ensure only admins can access this route
    if not current_user.is_admin:
        flash('Access denied. Only admins can generate question papers.', category='error')
        return redirect(url_for('auth.login'))

    current_year = datetime.now().year
    subjects = Subject.query.all()

    form_data = {
        'university_name': '',
        'exam_type': '',
        'academic_year': '',
        'exam_date': '',
        'start_time': '',
        'duration': '',
        'max_marks': '',
        'subject_code': '',
    }
    questions_by_module = None

    if request.method == 'POST':
        # Get form data
        form_data['exam_type'] = request.form.get('exam_type')
        form_data['academic_year'] = request.form.get('academic_year')
        form_data['exam_date'] = request.form.get('exam_date')
        form_data['start_time'] = request.form.get('start_time')
        form_data['duration'] = request.form.get('duration')
        form_data['max_marks'] = request.form.get('max_marks')
        form_data['subject_code'] = request.form.get('subject')
        

        # Validate required fields
        if not all([ form_data['exam_type'], form_data['academic_year'], form_data['exam_date'], form_data['start_time'], form_data['duration'], form_data['max_marks']]):
            flash('All fields are required. Please fill out the form completely.', category='error')
            return redirect(url_for('views.manual_generate_qp'))
        elif not form_data['subject_code']:
            flash('Please select a subject to proceed.', category='error')
        else:
            # Fetch the selected subject
            subject = Subject.query.filter_by(code=form_data['subject_code']).first()
            if not subject:
                flash('Invalid subject selected.', category='error')
            else:
                # Fetch module-wise questions
                questions_by_module = {}
                for module in range(1, 6):  # Assuming 5 modules
                    questions = Question.query.filter_by(
                        subject_code=form_data['subject_code'],
                        module_no=module,
                        is_verified=True  # Filter only verified questions
                    ).all()
                    questions_by_module[module] = questions

                action = request.form.get('action') 

                # Get the output format (PDF or Word)
                output_format = request.form.get('output_format')

                if action == 'manual-generate-qp' and output_format:
                    # Get selected question IDs from the form
                    selected_question_ids = request.form.getlist('selected_questions')

                    if not selected_question_ids:
                        flash('No questions selected. Please select questions to generate the question paper.', category='error')
                    else:
                        # Fetch selected questions
                        selected_questions = Question.query.filter(Question.id.in_(selected_question_ids)).all()

                        # Calculate total marks of selected questions
                        total_marks = sum(question.marks for question in selected_questions)

                        if total_marks != int(form_data['max_marks']):
                            flash(f'Total marks of selected questions ({total_marks}) do not match the required marks ({form_data["max_marks"]}). Please adjust your selection.', category='error')
                        else:
                            # Organize selected questions by module
                            selected_questions_by_module = {module: [] for module in range(1, 6)}
                            for question in selected_questions:
                                selected_questions_by_module[question.module_no].append(question)

                            # Render the HTML template for question paper
                            html = render_template(
                                'question_paper.html',
                                university_name=form_data['university_name'],
                                exam_type=form_data['exam_type'],
                                academic_year=form_data['academic_year'],
                                exam_date=form_data['exam_date'],
                                max_marks=form_data['max_marks'],
                                start_time=form_data['start_time'],
                                duration=form_data['duration'],
                                subject=subject,
                                department=subject.department,
                                questions_by_module=selected_questions_by_module,
                                additional_details=form_data
                            )

                            if output_format == 'pdf':
                                # Convert the HTML to a PDF
                                pdf = HTML(string=html, base_url=request.url_root).write_pdf()

                                # Create a BytesIO object for the binary content
                                response = BytesIO()
                                response.write(pdf)
                                response.seek(0)

                                return send_file(
                                    response,
                                    download_name="QuestionPaper.pdf",
                                    as_attachment=True
                                )
                               

                            elif output_format == 'word':
                                # Create Document
                                doc = Document()

                                # Add Logo (adjust path to the actual location)
                                image_path = os.path.join(os.getcwd(), 'website', 'static', 'images', 'uni.png')
                                doc.add_picture(image_path, width=Inches(6))
                                if doc.paragraphs:
                                    last_paragraph = doc.paragraphs[-1]
                                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


                                # Add Heading
                                heading = doc.add_paragraph('Department of Computer Science and Engineering')
                                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                heading.runs[0].bold = True
                                heading.runs[0].font.size = Pt(14)

                                # Add Exam Info
                                doc.add_paragraph(f"{form_data['exam_type']} – {form_data['academic_year']}", style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                doc.add_paragraph(f"Course Code: {subject.code} \t\t Maximum Marks: {form_data['max_marks']}")
                                doc.add_paragraph(f"Date: {form_data['exam_date']} \t\t Time: {form_data['start_time']} - Duration: {form_data['duration']} hours")

                                # Add Note
                                note = doc.add_paragraph('NOTE: ')
                                note.add_run('Answer All Questions').bold = True

                                # Create Table for Questions
                                table = doc.add_table(rows=1, cols=5)
                                table.style = 'Table Grid'

                                table.columns[0].width = Cm(0.001)  
                                table.columns[1].width = Cm(20) 
                                table.columns[2].width = Cm(0.001)  
                                table.columns[3].width = Cm(0.001)  
                                table.columns[4].width = Cm(0.001) 

                                headers = ['Q#', 'Question', 'Marks', 'BL', 'CO']
                                for i, header in enumerate(headers):
                                    cell = table.cell(0, i)
                                    cell.text = header
                                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    cell.paragraphs[0].runs[0].bold = True

                                # Add Questions to Table
                                question_number = 1
                                for module, questions in selected_questions_by_module.items():
                                    for question in questions:
                                        row_cells = table.add_row().cells
                                        row_cells[0].text = str(question_number)  # Q#
                                        row_cells[1].text = question.content  # Question
                                        row_cells[2].text = str(question.marks)  # Marks
                                        row_cells[3].text = question.rbt_level  # BL (Bloom's Level)
                                        row_cells[4].text = str(question.co_id)  # CO (Course Outcome)
                                        question_number += 1
                               
                                # Prepared and Approved Section
                                doc.add_paragraph("\nPrepared By: Dr Loganathan R\t\tApproved By: The Principal")

                                # Save to Memory Buffer
                                buffer = BytesIO()
                                doc.save(buffer)
                                buffer.seek(0)

                                # Send as Response
                                return send_file(
                                    buffer,
                                    download_name=f"{subject.name}_QuestionPaper.docx",
                                    as_attachment=True,
                                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                                )

    # Render the form with prefilled data
    return render_template(
        "admin/manual_generate_qp.html",
        questions_by_module=questions_by_module,
        subjects=subjects,
        current_year=current_year,
        form_data=form_data,
        user=current_user
    )